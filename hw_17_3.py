"""Только для оценки 10\10 в pylint"""
from docx import Document

if __name__ == "__main__":
    doc = Document()
    doc.add_paragraph("Hello Python")
    doc.save("hello_python.docx")

    doc = Document("hello_python.docx")
    bold_text = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bold_text += run.text

    print("Жирный текст из файла:", bold_text)

    doc_new = Document()
    paragraph = doc_new.add_paragraph("Это абзац с текстом.")
    run = paragraph.add_run()
    run.bold = True
    run.font.size = 20
    run.font.name = 'Arial'
    doc_new.save("new_document.docx")
